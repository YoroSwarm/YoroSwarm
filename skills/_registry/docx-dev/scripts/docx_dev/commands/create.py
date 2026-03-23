"""
create command - Create a new DOCX document from scratch.
Ported from C# CreateCommand.cs
"""
import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# Font configurations by document type
FONT_CONFIGS = {
    'report': {
        'body_font': 'Calibri',
        'heading_font': 'Calibri Light',
        'body_size': 11,
        'heading1_size': 26,
        'heading2_size': 22,
        'heading3_size': 18,
    },
    'letter': {
        'body_font': 'Calibri',
        'heading_font': 'Calibri Light',
        'body_size': 11,
        'heading1_size': 24,
        'heading2_size': 20,
        'heading3_size': 16,
    },
    'memo': {
        'body_font': 'Calibri',
        'heading_font': 'Calibri Light',
        'body_size': 11,
        'heading1_size': 22,
        'heading2_size': 18,
        'heading3_size': 14,
    },
    'academic': {
        'body_font': 'Times New Roman',
        'heading_font': 'Times New Roman',
        'body_size': 12,
        'heading1_size': 24,
        'heading2_size': 20,
        'heading3_size': 18,
    },
}

# Page sizes in twips (1 inch = 1440 twips)
PAGE_SIZES = {
    'letter': (12240, 15840),  # 8.5 x 11 inches
    'a4': (11906, 16838),
    'legal': (12240, 20160),  # 8.5 x 14 inches
    'a3': (16838, 23811),
}

# Margins in twips
MARGINS = {
    'standard': (1440, 1440, 1440, 1440),  # top, right, bottom, left
    'narrow': (720, 720, 720, 720),
    'wide': (1440, 1800, 1440, 1800),
}


def create_command(subparsers):
    parser = subparsers.add_parser('create', help='Create a new DOCX document from scratch')
    parser.add_argument('--output', required=True, help='Output DOCX file path')
    parser.add_argument('--type', default='report', help='Document type: report, letter, memo, academic')
    parser.add_argument('--title', help='Document title')
    parser.add_argument('--author', help='Document author')
    parser.add_argument('--page-size', default='letter', help='Page size: letter, a4, legal, a3')
    parser.add_argument('--margins', default='standard', help='Margin preset: standard, narrow, wide')
    parser.add_argument('--header', help='Header text')
    parser.add_argument('--footer', help='Footer text')
    parser.add_argument('--page-numbers', action='store_true', help='Add page numbers in footer')
    parser.add_argument('--toc', action='store_true', help='Insert table of contents placeholder')
    parser.add_argument('--content-json', help='Path to JSON file describing document content')
    parser.set_defaults(handler=handle_create)


def handle_create(args):
    """Main handler for create command."""
    doc = Document()

    # Set page size
    page_width, page_height = PAGE_SIZES.get(args.page_size.lower(), PAGE_SIZES['letter'])
    section = doc.sections[0]
    section.page_width = page_width
    section.page_height = page_height

    # Set margins
    top, right, bottom, left = MARGINS.get(args.margs.lower() if hasattr(args, 'margs') else 'standard', MARGINS['standard'])
    section.top_margin = top
    section.right_margin = right
    section.bottom_margin = bottom
    section.left_margin = left

    # Configure header
    if args.header:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.text = args.header
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Configure footer
    if args.footer or args.page_numbers:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]

        if args.footer:
            p.text = args.footer
            if args.page_numbers:
                p.text += " — "
                p.add_run()

        if args.page_numbers:
            # Add page number field
            run = p.add_run()
            fldChar1 = run._r.append_element(qn('w:fldChar'))
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = run._r.append_element(qn('w:instrText'))
            instrText.text = ' PAGE '

            fldChar2 = run._r.append_element(qn('w:fldChar'))
            fldChar2.set(qn('w:fldCharType'), 'end')

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Get font config
    config = FONT_CONFIGS.get(args.type.lower(), FONT_CONFIGS['report'])

    # Title
    if args.title:
        title = doc.add_heading(args.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = config['heading_font']
            run.font.size = Pt(config['heading1_size'])

    # Author subtitle
    if args.author:
        subtitle = doc.add_paragraph(args.author)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.name = config['body_font']
        subtitle_run.font.size = Pt(config['body_size'] + 2)
        subtitle_run.font.color.rgb = None

    # TOC placeholder
    if args.toc:
        doc.add_heading('Table of Contents', level=1)
        toc_para = doc.add_paragraph('Update this field to generate table of contents.')
        # Add page break
        doc.add_page_break()

    # Content from JSON
    if args.content_json:
        content_path = Path(args.content_json)
        if content_path.exists():
            with open(content_path, 'r', encoding='utf-8') as f:
                content_data = json.load(f)
            add_content_from_json(doc, content_data, config)

    # Save
    doc.save(args.output)
    print(f"Created {args.type} document: {args.output}")
    return 0


def add_content_from_json(doc, content_data, config):
    """Add content from JSON to document."""
    if not isinstance(content_data, list):
        return

    for item in content_data:
        item_type = item.get('type', 'paragraph')
        text = item.get('text', '')

        if item_type == 'heading':
            level = min(max(item.get('level', 1), 1), 6)
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.name = config['heading_font']

        elif item_type == 'paragraph':
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.name = config['body_font']
                run.font.size = Pt(config['body_size'])

        elif item_type == 'pagebreak':
            doc.add_page_break()
