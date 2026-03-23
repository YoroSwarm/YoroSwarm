"""
edit command - Edit content in existing DOCX.
Ported from C# EditContentCommand.cs
"""
import argparse
import json
import sys
from pathlib import Path
from zipfile import ZipFile
from lxml import etree
from docx import Document

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def edit_command(subparsers):
    parser = subparsers.add_parser('edit', help='Edit content in existing DOCX')
    subparsers.required = True

    sub = parser.add_subparsers(dest='edit_action', help='Edit actions')

    # replace-text
    replace_parser = sub.add_parser('replace-text', help='Replace text in document')
    replace_parser.add_argument('--input', required=True, help='Input DOCX file')
    replace_parser.add_argument('--output', required=True, help='Output DOCX file')
    replace_parser.add_argument('--find', required=True, help='Text to find')
    replace_parser.add_argument('--replace', required=True, help='Text to replace with')
    replace_parser.set_defaults(handler=handle_replace_text)

    # fill-placeholders
    fill_parser = sub.add_parser('fill-placeholders', help='Fill placeholders from JSON data')
    fill_parser.add_argument('--input', required=True, help='Input DOCX file')
    fill_parser.add_argument('--output', required=True, help='Output DOCX file')
    fill_parser.add_argument('--data', required=True, help='JSON data to fill')
    fill_parser.set_defaults(handler=handle_fill_placeholders)

    # fill-table
    fill_table_parser = sub.add_parser('fill-table', help='Fill table cells from JSON')
    fill_table_parser.add_argument('--input', required=True, help='Input DOCX file')
    fill_table_parser.add_argument('--output', required=True, help='Output DOCX file')
    fill_table_parser.add_argument('--data', required=True, help='JSON data for table')
    fill_table_parser.set_defaults(handler=handle_fill_table)

    # insert-section
    insert_parser = sub.add_parser('insert-section', help='Insert a new section')
    insert_parser.add_argument('--input', required=True, help='Input DOCX file')
    insert_parser.add_argument('--output', required=True, help='Output DOCX file')
    insert_parser.add_argument('--title', required=True, help='Section title')
    insert_parser.add_argument('--content', default='', help='Section content')
    insert_parser.set_defaults(handler=handle_insert_section)

    # remove-section
    remove_parser = sub.add_parser('remove-section', help='Remove a section by title')
    remove_parser.add_argument('--input', required=True, help='Input DOCX file')
    remove_parser.add_argument('--output', required=True, help='Output DOCX file')
    remove_parser.add_argument('--title', required=True, help='Section title to remove')
    remove_parser.set_defaults(handler=handle_remove_section)

    # update-header-footer
    header_footer_parser = sub.add_parser('update-header-footer', help='Update header/footer content')
    header_footer_parser.add_argument('--input', required=True, help='Input DOCX file')
    header_footer_parser.add_argument('--output', required=True, help='Output DOCX file')
    header_footer_parser.add_argument('--header', help='New header text')
    header_footer_parser.add_argument('--footer', help='New footer text')
    header_footer_parser.set_defaults(handler=handle_update_header_footer)


def handle_replace_text(args):
    """Replace text in document."""
    try:
        doc = Document(args.input)

        replacements = 0
        for para in doc.paragraphs:
            if args.find in para.text:
                para.text = para.text.replace(args.find, args.replace)
                replacements += 1

            for run in para.runs:
                if args.find in run.text:
                    run.text = run.text.replace(args.find, args.replace)
                    replacements += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if args.find in para.text:
                            para.text = para.text.replace(args.find, args.replace)
                            replacements += 1

        doc.save(args.output)
        print(f"Replaced {replacements} occurrences")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


def handle_fill_placeholders(args):
    """Fill placeholders with data from JSON."""
    try:
        with open(args.data, 'r', encoding='utf-8') as f:
            data = json.load(f)

        doc = Document(args.input)

        for key, value in data.items():
            placeholder = f"{{{key}}}"

            for para in doc.paragraphs:
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))

                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if placeholder in para.text:
                                para.text = para.text.replace(placeholder, str(value))

        doc.save(args.output)
        print(f"Filled {len(data)} placeholders")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


def handle_fill_table(args):
    """Fill table cells with JSON data."""
    try:
        with open(args.data, 'r', encoding='utf-8') as f:
            data = json.load(f)

        doc = Document(args.input)

        if not data or not isinstance(data, list):
            print("Error: data must be a list of row objects", file=sys.stderr)
            return 1

        table_index = 0
        for table in doc.tables:
            if table_index >= len(data):
                break

            row_data = data[table_index]
            if not isinstance(row_data, dict):
                table_index += 1
                continue

            headers = [cell.text.strip() for cell in table.rows[0].cells]

            for row_idx, row in enumerate(table.rows[1:], 1):
                if row_idx - 1 >= len(data):
                    break

                row_data = data[row_idx - 1]
                if not isinstance(row_data, dict):
                    continue

                for col_idx, cell in enumerate(row.cells):
                    if col_idx < len(headers):
                        header = headers[col_idx]
                        if header in row_data:
                            cell.text = str(row_data[header])

            table_index += 1

        doc.save(args.output)
        print(f"Filled table at index {table_index}")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


def handle_insert_section(args):
    """Insert a new section."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(args.input)

        # Add section heading
        heading = doc.add_heading(args.title, level=1)

        # Add content
        if args.content:
            doc.add_paragraph(args.content)

        doc.save(args.output)
        print(f"Inserted section: {args.title}")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


def handle_remove_section(args):
    """Remove a section by title."""
    try:
        doc = Document(args.input)

        removed = False
        paragraphs_to_remove = []

        for para in doc.paragraphs:
            if para.text.strip() == args.title.strip():
                paragraphs_to_remove.append(para)
                removed = True

        for para in paragraphs_to_remove:
            p = para._element
            p.getparent().remove(p)

        doc.save(args.output)
        print(f"Removed section: {args.title}" if removed else "Section not found")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


def handle_update_header_footer(args):
    """Update header and footer content."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(args.input)

        for section in doc.sections:
            if args.header:
                header = section.header
                header.is_linked_to_previous = False
                for para in header.paragraphs:
                    para.clear()
                header.paragraphs[0].text = args.header
                header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            if args.footer:
                footer = section.footer
                footer.is_linked_to_previous = False
                for para in footer.paragraphs:
                    para.clear()
                footer.paragraphs[0].text = args.footer
                footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(args.output)
        print("Updated header/footer")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
